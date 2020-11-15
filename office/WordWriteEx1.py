from docx import Document
from docx.shared import Inches

document = Document()

#添加标题，并设置级别，范围：0 至 9，默认为1
document.add_heading('Document Title', 0)

#添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
p = document.add_paragraph('A plain paragraph having some ')
#在段落后面追加文本，并可设置样式
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

#添加项目列表（前面一个小圆点）
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph('second item in unordered list', style='List Bullet')

#添加项目列表（前面数字）
document.add_paragraph('first item in ordered list', style='List Number')
document.add_paragraph('second item in ordered list', style='List Number')

#添加图片
document.add_picture("G:\projects\PYTHON-python-study\\file\\test.jpg" ,width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

#添加表格：一行三列
# 表格样式参数可选：
# Normal Table
# Table Grid
# Light Shading、 Light Shading Accent 1 至 Light Shading Accent 6
# Light List、Light List Accent 1 至 Light List Accent 6
# Light Grid、Light Grid Accent 1 至 Light Grid Accent 6
# 太多了其它省略...
table = document.add_table(rows=1, cols=3, style='Light Shading Accent 2')
#获取第一行的单元格列表
hdr_cells = table.rows[0].cells
#下面三行设置上面第一行的三个单元格的文本值
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    #表格添加行，并返回行所在的单元格列表
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

#保存.docx文档
document.save('demo.docx')